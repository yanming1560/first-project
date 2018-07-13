from docx import Document
from docx.shared import Cm

file=Document()
file.add_heading('Title',0)
file.add_heading('first',1)
file.add_heading('f.1',2)
file.add_paragraph('hahahahhahahahahha1')
file.add_heading('f.2',2)
file.add_paragraph('hahahahhahahahahha2')
file.add_heading('second',1)
file.add_paragraph('oooooooooo')
secs=file.sections[0]
secs.left_margin=Cm(1.5)
secs.right_margin=Cm(1.5)
secs.top_margin=Cm(1.5)
secs.bottom_margin=Cm(1.5)
file.save('test.docx')
